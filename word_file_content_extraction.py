import os
import docx
import pytesseract
from PIL import Image
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn

# Set Tesseract path if needed (Uncomment for Windows users)
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def extract_text(doc_path):
    """ Extract all text from a Word file. """
    doc = docx.Document(doc_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_tables(doc_path):
    """ Extract tables as lists of lists from a Word file. """
    doc = docx.Document(doc_path)
    tables_data = []
    
    for table in doc.tables:
        table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables_data.append(table_data)
    
    return tables_data

def extract_images(doc_path, output_folder="images"):
    """ Extract and save images from a Word file. """
    doc = docx.Document(doc_path)
    
    # Ensure output directory exists
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    image_files = []
    for rel in doc.part.rels:
        if doc.part.rels[rel].reltype == RELATIONSHIP_TYPE.IMAGE:
            image_part = doc.part.rels[rel].target_part
            image_data = image_part.blob
            image_ext = image_part.content_type.split("/")[-1]  # Get extension
            
            image_filename = os.path.join(output_folder, f"image_{len(image_files) + 1}.{image_ext}")
            with open(image_filename, "wb") as img_file:
                img_file.write(image_data)
            
            image_files.append(image_filename)

    return image_files

def extract_links(doc_path):
    """ Extract hyperlinks from a Word file. """
    doc = docx.Document(doc_path)
    links = []

    for rel in doc.part.rels:
        if doc.part.rels[rel].reltype == RELATIONSHIP_TYPE.HYPERLINK:
            link = doc.part.rels[rel].target_ref
            links.append(link)

    return links

def extract_text_from_images(image_files):
    """ Perform OCR on extracted images to get text. """
    extracted_texts = {}
    
    for img_path in image_files:
        text = pytesseract.image_to_string(Image.open(img_path))
        extracted_texts[img_path] = text.strip()
    
    return extracted_texts


# ======== Example Usage ========

doc_path = "sample_doc.docx"  # Replace with your Word file
output_folder = "extracted_images"

# Extract text
text = extract_text(doc_path)
print("Extracted Text:\n", text, "\n")

# Extract tables
tables = extract_tables(doc_path)
for i, table in enumerate(tables, 1):
    print(f"Table {i}:\n", table, "\n")

# Extract images
image_files = extract_images(doc_path, output_folder)
print("Extracted Images:", image_files)

# Extract links
links = extract_links(doc_path)
print("Extracted Links:")
for link in links:
    print(link)

# Extract text from images using OCR
ocr_texts = extract_text_from_images(image_files)
print("OCR Extracted Text from Images:\n", ocr_texts)
